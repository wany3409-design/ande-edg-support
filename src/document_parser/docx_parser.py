"""
DOCX 文档解析器

基于 python-docx 提取文本，支持：
- 提取段落文本
- 保留标题/章节信息
- 提取表格文本
"""

from typing import List

from docx import Document

from src.document_parser.base_parser import BaseDocumentParser
from src.document_parser.models import RawDocument, ParsedPage


class DocxParser(BaseDocumentParser):
    """DOCX 文档解析器"""

    supported_extensions = [".docx"]
    parser_name = "docx"

    def _do_parse(self, raw_doc: RawDocument) -> List[ParsedPage]:
        try:
            doc = Document(raw_doc.file_path)
        except Exception as e:
            raise RuntimeError(f"无法打开DOCX文件: {raw_doc.file_path}, 错误: {e}")

        current_section = ""
        all_text_parts = []
        page_estimate = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 识别标题作为章节
            if para.style and para.style.name and "Heading" in para.style.name:
                current_section = text

            all_text_parts.append(text)

            # 按段落数粗略估算分页（每~30段落算一页）
            if len(all_text_parts) % 30 == 0:
                page_estimate += 1

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    all_text_parts.append(row_text)

        # 组装为伪页面
        text_by_pages = self._split_into_pages(all_text_parts, page_estimate)
        pages = [
            ParsedPage(page_num=i + 1, text=t, section=current_section)
            for i, t in enumerate(text_by_pages)
        ]

        return pages

    def _split_into_pages(self, parts: List[str], page_count: int) -> List[str]:
        """将段落列表按页数分割"""
        if not parts:
            return [""]

        page_size = max(1, len(parts) // max(1, page_count))
        result = []
        for i in range(0, len(parts), page_size):
            chunk = "\n".join(parts[i:i + page_size])
            if chunk.strip():
                result.append(chunk)
        return result if result else ["\n".join(parts)]
